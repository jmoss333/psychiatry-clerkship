#!/usr/bin/env python3
"""
OpenEvidence folder scanner for the clerkship library.

Detects NEW or CHANGED review files dropped into the
"OPENEVIDENCE RAW FILES TO REVIEW" folder (compared to a manifest of
already-processed files), extracts their text to a staging area, and can
mark files as processed once reviewed.

Path resolution is location-relative (works regardless of session mount
prefix): the scanned folder is <library_root>/"OPENEVIDENCE RAW FILES TO REVIEW",
where <library_root> is four directories up from this script
(.../13_Faculty_Resources/_automation/oe_scanner/oe_scan.py).

Usage:
  python3 oe_scan.py                 # report new/changed files + extract their text to ./staging/
  python3 oe_scan.py --commit "A.docx" "B.docx"   # mark specific files processed (record current hash)
  python3 oe_scan.py --commit-all    # mark every currently-present file processed (use to seed)
  python3 oe_scan.py --list          # just list tracked/untracked without extracting
  python3 oe_scan.py --pending       # just list teaching pages still carrying a "Pending re-attestation" tag
  python3 oe_scan.py --attest "<page>" ["<page>" ...]   # sign off: clear the tag, restamp the page, log it
  python3 oe_scan.py --attest-all    # attest every page currently in the pending queue
  python3 oe_scan.py --attest-log    # print the attestation audit trail from the manifest

The default scan output also includes a `pending_attestation` list: teaching pages
where new clinical facts were inserted (via this review pipeline) but Dr. Moss has
not yet signed off. This turns the attestation queue into an automated to-do that
surfaces in every weekly run's summary.

Clearing an item is what `--attest` does, and it is the only supported way to do it.
For each named page it (1) strips the "Pending re-attestation" tag from the review-status
line, (2) advances the "Reviewed and attested by ... (DATE)" stamp to today, and
(3) appends a record to the manifest's `__attestations__` audit log capturing the note
that was cleared, the prior stamp date, and the attester — so the sign-off survives even
though the tag is gone from the page. Pair with --dry-run to preview the exact rewrite
before touching anything.

Flags usable alongside --attest / --attest-all:
  --dry-run                 show the before/after line for each page, write nothing
  --attester "Name, MD"     override the attester (default: Joshua Moss, MD)
  --date YYYY-MM-DD         override the stamp date (default: today)
"""
import os, sys, json, hashlib, datetime, re

ATTEST_LOG_KEY = "__attestations__"   # reserved manifest key; cannot collide with a filename
RESERVED_KEYS = {ATTEST_LOG_KEY}
DEFAULT_ATTESTER = "Joshua Moss, MD"

HERE = os.path.dirname(os.path.abspath(__file__))
LIB_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(HERE)))  # up: oe_scanner -> _automation -> 13_Faculty_Resources -> <lib>
FOLDER = os.path.join(LIB_ROOT, "OPENEVIDENCE RAW FILES TO REVIEW")
MANIFEST = os.path.join(HERE, "oe_manifest.json")
STAGING = os.path.join(HERE, "staging")
EXTS = {".docx", ".pdf", ".txt", ".md", ".csv"}

# --- Pending-attestation queue (read-only scan of teaching pages) ---
PENDING_MARK = "pending re-attestation"
# path segments to skip: input/automation/archive/build dirs, not learner-facing teaching pages
_SKIP_SEGMENTS = {"99_Archive", "Handoffs", "node_modules", "outputs", "tmp", "tests",
                  "staging", "docs", "sp-proxy", "faculty-console", "quick-wins",
                  "00_START_HERE", "OPENEVIDENCE RAW FILES TO REVIEW",
                  "Evidence Inbox", "_inbox"}

def _sha(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()

def _load():
    try:
        return json.load(open(MANIFEST, encoding="utf-8"))
    except Exception:
        return {}

def _save(m):
    json.dump(m, open(MANIFEST, "w", encoding="utf-8"), indent=2, sort_keys=True)

def _candidates():
    out = []
    if not os.path.isdir(FOLDER):
        return out
    for fn in sorted(os.listdir(FOLDER)):
        p = os.path.join(FOLDER, fn)
        if not os.path.isfile(p):
            continue
        if fn.startswith("~$") or fn.startswith("."):
            continue
        if fn.lower() == "readme.md":     # a drop folder's own instructions are not a drop
            continue
        if os.path.splitext(fn)[1].lower() not in EXTS:
            continue
        out.append(fn)
    return out

def _new_or_changed():
    m = _load()
    res = []
    for fn in _candidates():
        p = os.path.join(FOLDER, fn)
        s = _sha(p)
        if m.get(fn, {}).get("sha256") != s:
            res.append((fn, s))
    return m, res

def _extract(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            try:
                import docx  # python-docx
            except Exception:
                os.system("pip install python-docx --break-system-packages -q >/dev/null 2>&1")
                import docx
            d = docx.Document(path)
            parts = [x.text for x in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)
        if ext in (".txt", ".md", ".csv"):
            return open(path, encoding="utf-8", errors="ignore").read()
        if ext == ".pdf":
            try:
                import pypdf
            except Exception:
                os.system("pip install pypdf --break-system-packages -q >/dev/null 2>&1")
                import pypdf
            r = pypdf.PdfReader(path)
            return "\n".join((pg.extract_text() or "") for pg in r.pages)
    except Exception as e:
        return "[extraction error: %s]" % e
    return ""

def _scan_pending_attestation():
    """Walk the library for teaching pages still carrying a 'Pending re-attestation'
    tag (added when a new clinical fact is inserted before Dr. Moss re-attests).
    Read-only and additive; surfaces the attestation queue in every scan summary."""
    hits = []
    pat = re.compile(r"pending re-attestation:?\**\s*(.*)", re.IGNORECASE)
    for root, dirs, files in os.walk(LIB_ROOT):
        # prune input/automation/archive/build dirs and any dot/underscore dirs
        dirs[:] = [d for d in dirs
                   if not d.startswith(".") and not d.startswith("_")
                   and d not in _SKIP_SEGMENTS]
        for fn in files:
            if not fn.lower().endswith(".md"):
                continue
            p = os.path.join(root, fn)
            try:
                with open(p, encoding="utf-8", errors="ignore") as fh:
                    for i, line in enumerate(fh, 1):
                        if PENDING_MARK in line.lower():
                            mo = pat.search(line)
                            note = (mo.group(1) if mo else line).strip().rstrip(".")
                            hits.append({"page": os.path.relpath(p, LIB_ROOT),
                                         "line": i, "note": note})
            except Exception:
                continue
    return sorted(hits, key=lambda h: h["page"])

# --- Attestation (the only supported way to clear a pending item) ---
# Matches the tag plus any leading separator, from the separator to end of line:
#   " · **Pending re-attestation:** CATIE discontinuation figure ... added 2026-07-23."
_PENDING_TAG_RE = re.compile(
    r"(?:\s*[·—–|-]\s*)?\*{0,2}Pending re-attestation\s*:?\*{0,2}.*$",
    re.IGNORECASE)
# Matches the sign-off stamp so the date (and optionally the name) can be advanced:
#   "Reviewed and attested by Joshua Moss, MD (2026-07-09)"
_STAMP_RE = re.compile(r"(Reviewed and attested by\s+)(.+?)\s*\((\d{4}-\d{2}-\d{2})\)")


def _resolve_page(arg):
    """Accept a LIB_ROOT-relative path, an absolute path, or a bare filename."""
    cand = arg if os.path.isabs(arg) else os.path.join(LIB_ROOT, arg)
    if os.path.isfile(cand):
        return cand
    base = os.path.basename(arg).lower()          # bare-filename fallback
    matches = [h["page"] for h in _scan_pending_attestation()
               if os.path.basename(h["page"]).lower() == base]
    if len(matches) == 1:
        return os.path.join(LIB_ROOT, matches[0])
    return None


def _attest_page(path, attester, date_str, dry_run=False):
    """Strip the pending tag, advance the sign-off stamp, return an audit record.
    Rewrites only the review-status line(s); all other page content is untouched."""
    rel = os.path.relpath(path, LIB_ROOT)
    with open(path, encoding="utf-8") as fh:
        lines = fh.readlines()

    cleared, edits, prior_date, stamp_found = [], [], None, False
    for i, line in enumerate(lines):
        if PENDING_MARK not in line.lower():
            continue
        before = line.rstrip("\n")
        mo = re.search(r"Pending re-attestation\s*:?\**\s*(.*?)\s*$", before, re.IGNORECASE)
        cleared.append(mo.group(1).strip().rstrip(".") if mo else before.strip())

        after = _PENDING_TAG_RE.sub("", before).rstrip()
        sm = _STAMP_RE.search(after)
        if sm:
            stamp_found = True
            prior_date = sm.group(3)
            after = _STAMP_RE.sub(
                lambda m: "%s%s (%s)" % (m.group(1), attester, date_str), after, count=1)
        lines[i] = after + "\n"
        edits.append({"line": i + 1, "before": before, "after": after})

    if not edits:
        return {"page": rel, "status": "no_pending_tag"}
    if not dry_run:
        with open(path, "w", encoding="utf-8") as fh:
            fh.writelines(lines)
    rec = {"page": rel, "status": "would_attest" if dry_run else "attested",
           "cleared_notes": cleared, "prior_attestation": prior_date,
           "new_attestation": date_str, "attester": attester, "edits": edits}
    if not stamp_found:
        rec["warning"] = ("no 'Reviewed and attested by ... (YYYY-MM-DD)' stamp found on the "
                          "tagged line; tag cleared but no date advanced — check this page by hand")
    return rec


def cmd_attest(pages, attester=DEFAULT_ATTESTER, date_str=None, dry_run=False, all_pending=False):
    date_str = date_str or datetime.date.today().isoformat()
    if all_pending:
        pages = sorted({h["page"] for h in _scan_pending_attestation()})
    out = {"ran_at": datetime.datetime.now().isoformat(timespec="seconds"),
           "attester": attester, "stamp_date": date_str, "dry_run": dry_run,
           "requested": len(pages), "results": []}
    if not pages:
        out["note"] = "nothing to attest — pending queue is empty"
        print(json.dumps(out, indent=2))
        return

    m = _load()
    log = m.get(ATTEST_LOG_KEY, [])
    wrote = False
    for p in pages:
        path = _resolve_page(p)
        if not path:
            out["results"].append({"page": p, "status": "not_found"})
            continue
        rec = _attest_page(path, attester, date_str, dry_run)
        out["results"].append(rec)
        if rec["status"] == "attested":
            wrote = True
            log.append({"page": rec["page"], "attested_at": out["ran_at"],
                        "attester": attester, "stamp_date": date_str,
                        "prior_attestation": rec.get("prior_attestation"),
                        "cleared_notes": rec.get("cleared_notes", [])})
    if wrote:
        m[ATTEST_LOG_KEY] = log
        _save(m)
    out["attested_count"] = sum(1 for r in out["results"] if r["status"] == "attested")
    out["remaining_pending"] = len(_scan_pending_attestation())
    print(json.dumps(out, indent=2))


def cmd_attest_log():
    print(json.dumps({"attestations": _load().get(ATTEST_LOG_KEY, [])}, indent=2))


def cmd_report(extract=True):
    m, new = _new_or_changed()
    os.makedirs(STAGING, exist_ok=True)
    pending = _scan_pending_attestation()
    report = {
        "scanned_at": datetime.datetime.now().isoformat(timespec="seconds"),
        "folder": FOLDER,
        "total_files": len(_candidates()),
        "already_processed": len([k for k in m if k not in RESERVED_KEYS]),
        "new_or_changed_count": len(new),
        "pending_attestation_count": len(pending),
        "new_or_changed": [],
        "pending_attestation": pending,
    }
    for fn, s in new:
        entry = {"file": fn, "sha256": s}
        if extract:
            txt = _extract(os.path.join(FOLDER, fn))
            stage = os.path.join(STAGING, fn + ".txt")
            open(stage, "w", encoding="utf-8").write(txt)
            entry["chars"] = len(txt)
            entry["staged_text"] = stage
        report["new_or_changed"].append(entry)
    print(json.dumps(report, indent=2))

def cmd_commit(files):
    m = _load()
    present = set(_candidates())
    done = []
    if files == ["--all"]:
        files = sorted(present)
    for fn in files:
        if fn not in present:
            continue
        m[fn] = {"sha256": _sha(os.path.join(FOLDER, fn)),
                 "processed_at": datetime.datetime.now().isoformat(timespec="seconds")}
        done.append(fn)
    _save(m)
    print(json.dumps({"committed": done,
                      "manifest_size": len([k for k in m if k not in RESERVED_KEYS])}, indent=2))

def _pop_opts(args):
    """Pull --dry-run / --attester X / --date Y out of args; return (rest, opts)."""
    rest, opts = [], {"dry_run": False, "attester": DEFAULT_ATTESTER, "date_str": None}
    i = 0
    while i < len(args):
        a = args[i]
        if a == "--dry-run":
            opts["dry_run"] = True
        elif a == "--attester" and i + 1 < len(args):
            i += 1; opts["attester"] = args[i]
        elif a == "--date" and i + 1 < len(args):
            i += 1; opts["date_str"] = args[i]
        else:
            rest.append(a)
        i += 1
    if opts["date_str"] and not re.fullmatch(r"\d{4}-\d{2}-\d{2}", opts["date_str"]):
        sys.exit("--date must be YYYY-MM-DD, got: %s" % opts["date_str"])
    return rest, opts


def _pop_inbox_opts(args):
    """Retarget the scanner at a different drop folder / ledger.

    Defaults are unchanged, so every existing invocation and the OpenEvidence
    manifest behave exactly as before. This exists so one scanner can serve more
    than one inbox: the OpenEvidence review folder and the general evidence
    inbox each keep their own ledger.

        oe_scan.py --folder "<abs path>" --manifest "<abs path>" [--list|--commit ...]
    """
    global FOLDER, MANIFEST, STAGING
    rest = []
    i = 0
    while i < len(args):
        if args[i] == "--folder" and i + 1 < len(args):
            FOLDER = os.path.abspath(os.path.expanduser(args[i + 1])); i += 2
        elif args[i] == "--manifest" and i + 1 < len(args):
            MANIFEST = os.path.abspath(os.path.expanduser(args[i + 1]))
            STAGING = os.path.join(os.path.dirname(MANIFEST), "staging")
            i += 2
        else:
            rest.append(args[i]); i += 1
    return rest


if __name__ == "__main__":
    args = _pop_inbox_opts(sys.argv[1:])
    if not args:
        cmd_report(extract=True)
    elif args[0] == "--list":
        cmd_report(extract=False)
    elif args[0] == "--commit-all":
        cmd_commit(["--all"])
    elif args[0] == "--commit":
        cmd_commit(args[1:] if len(args) > 1 else [])
    elif args[0] == "--pending":
        print(json.dumps({"pending_attestation": _scan_pending_attestation()}, indent=2))
    elif args[0] == "--attest":
        rest, opts = _pop_opts(args[1:])
        if not rest:
            sys.exit('--attest needs at least one page path (or use --attest-all). '
                     'Run --pending to see the queue.')
        cmd_attest(rest, **opts)
    elif args[0] == "--attest-all":
        _, opts = _pop_opts(args[1:])
        cmd_attest([], all_pending=True, **opts)
    elif args[0] == "--attest-log":
        cmd_attest_log()
    else:
        print("unknown args:", args)
